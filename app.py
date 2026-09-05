
import streamlit as st
from docx import Document
from io import BytesIO
import requests


# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="AI Content Repurposer",
    page_icon="✨",
    layout="wide"
)


# ---------------- AI FUNCTION ----------------
def ask_ai(prompt):
    api_key = st.secrets["GROQ_API_KEY"]

    response = requests.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        },
        json={
            "model": "llama-3.1-8b-instant",
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "You are an expert social media content repurposing AI. "
                        "Use ONLY the information provided by the user. "
                        "Never invent facts, statistics, names, dates, links, "
                        "products, customers, results or events."
                    )
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.4
        },
        timeout=120
    )

    if response.status_code != 200:
        raise Exception(response.text)

    return response.json()["choices"][0]["message"]["content"]


# ---------------- FILE READER ----------------
def extract_file_content(uploaded_file):

    if uploaded_file.name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8")

    elif uploaded_file.name.endswith(".docx"):
        document = Document(uploaded_file)

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        return "\n".join(paragraphs)

    return ""


# ---------------- CONTENT ANALYSIS ----------------
def analyze_content(content):

    prompt = f"""
Analyze ONLY the ORIGINAL CONTENT.

ORIGINAL CONTENT:
-------------------------
{content}
-------------------------

Return exactly:

MAIN TOPIC:
[main topic]

MAIN MESSAGE:
[main message in one sentence]

KEY FACTS:
- [fact]
- [fact]
- [fact]

KEY POINTS:
- [point]
- [point]
- [point]

TARGET AUDIENCE:
[only if supported by the content]

TONE:
[tone]

IMPORTANT KEYWORDS:
- [keyword]
- [keyword]
- [keyword]

STRICT RULES:
- Use only information from the original content.
- Do not add outside facts.
- Do not invent statistics.
- Do not invent examples.
- If information is unavailable, write "Not specified".
"""

    return ask_ai(prompt)


# ---------------- GENERATION ----------------
def generate_platform_content(content, analysis, platform, tone):

    instructions = {

        "Instagram": """
Create:
CAPTION:
REEL HOOK:
REEL SCRIPT:
HASHTAGS:
CTA:
""",

        "LinkedIn": """
Create:
HOOK:
POST:
HASHTAGS:
CTA:
""",

        "X": """
Create:
POST:
THREAD:
HASHTAGS:
""",

        "YouTube": """
Create:
TITLE:
DESCRIPTION:
KEYWORDS/TAGS:
HOOK:
"""
    }

    prompt = f"""
You are repurposing ORIGINAL CONTENT for {platform}.

ORIGINAL CONTENT:
-------------------------
{content}
-------------------------

CONTENT ANALYSIS:
-------------------------
{analysis}
-------------------------

TONE:
{tone}

Create platform-specific content using this structure:

{instructions[platform]}

STRICT RULES:
- Original content is the source of truth.
- Use ONLY facts and ideas present in the original content.
- Do not invent statistics.
- Do not invent names.
- Do not invent dates.
- Do not invent links.
- Do not invent products or services.
- Do not invent customer results.
- Do not make unsupported claims.
- Preserve important terminology.
- Do not replace specific terms like AI with vague phrases such as
  "this technology", "new tools", or "intelligent systems".
- Make the content natural and publish-ready.
- Do not explain what you changed.
- Do not mention AI.
- Do not ask questions.
- Return ONLY the requested content.
"""

    return ask_ai(prompt)


# ---------------- VALIDATION ----------------
def validate_content(original_content, generated_content, platform):

    prompt = f"""
You are a strict content fact-checking and editing system.

ORIGINAL CONTENT:
-------------------------
{original_content}
-------------------------

GENERATED {platform.upper()} CONTENT:
-------------------------
{generated_content}
-------------------------

Check the generated content against the original.

Remove or rewrite anything that is NOT supported by the original content.

Unsupported information includes:
- Invented statistics
- Invented links
- Invented products
- Invented services
- Invented events
- Invented customer results
- Unsupported claims
- Information from outside knowledge

IMPORTANT:
- Preserve supported information.
- Preserve the platform structure.
- Do not add new information.
- Do not explain changes.
- Return ONLY the corrected final content.

If everything is supported, return the generated content unchanged.
"""

    return ask_ai(prompt)


# ---------------- DOCX EXPORT ----------------
def create_docx(results):

    document = Document()

    document.add_heading(
        "AI-Based Social Media Content Repurposing System",
        level=1
    )

    for platform, content in results.items():

        document.add_heading(platform, level=2)

        for line in content.split("\n"):

            if line.strip():
                document.add_paragraph(line)

    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output


# ---------------- TXT EXPORT ----------------
def create_txt(results):

    text = ""

    for platform, content in results.items():

        text += f"\n{'=' * 50}\n"
        text += f"{platform.upper()}\n"
        text += f"{'=' * 50}\n\n"
        text += content
        text += "\n"

    return text


# ---------------- HEADER ----------------
st.title("✨ AI-Based Social Media Content Repurposing System")

st.caption(
    "Transform one piece of content into platform-specific content "
    "for Instagram, LinkedIn, X and YouTube."
)


# ---------------- SIDEBAR ----------------
with st.sidebar:

    st.header("⚙️ Settings")

    tone = st.selectbox(
        "Select Tone",
        [
            "Professional",
            "Casual",
            "Creative",
            "Educational",
            "Friendly",
            "Persuasive"
        ]
    )

    st.markdown("---")

    st.write("Select platforms:")

    instagram = st.checkbox("📸 Instagram", value=True)
    linkedin = st.checkbox("💼 LinkedIn", value=True)
    x_platform = st.checkbox("𝕏 X", value=True)
    youtube = st.checkbox("▶️ YouTube", value=True)


# ---------------- INPUT ----------------
st.subheader("📝 Original Content")

input_method = st.radio(
    "Choose input method",
    ["Paste Content", "Upload TXT/DOCX"],
    horizontal=True
)

content = ""

if input_method == "Paste Content":

    content = st.text_area(
        "Paste your original content here:",
        height=250,
        placeholder="Paste your article, blog, script, transcript or other content..."
    )

else:

    uploaded_file = st.file_uploader(
        "Upload TXT or DOCX file",
        type=["txt", "docx"]
    )

    if uploaded_file:

        content = extract_file_content(uploaded_file)

        st.success("File loaded successfully.")

        with st.expander("Preview Original Content"):
            st.write(content)


# ---------------- GENERATE ----------------
if st.button(
    "🚀 Repurpose Content",
    type="primary",
    use_container_width=True
):

    if not content.strip():

        st.warning("Please provide some content first.")

    else:

        platforms = []

        if instagram:
            platforms.append("Instagram")

        if linkedin:
            platforms.append("LinkedIn")

        if x_platform:
            platforms.append("X")

        if youtube:
            platforms.append("YouTube")

        if not platforms:

            st.warning("Please select at least one platform.")

        else:

            try:

                # -------- ANALYSIS --------
                with st.spinner("🔎 Analyzing original content..."):

                    analysis = analyze_content(content)

                with st.expander("🧠 AI Content Analysis"):

                    st.text(analysis)

                # -------- GENERATION --------
                results = {}

                progress = st.progress(0)

                for i, platform in enumerate(platforms):

                    with st.spinner(
                        f"✍️ Creating {platform} content..."
                    ):

                        generated = generate_platform_content(
                            content,
                            analysis,
                            platform,
                            tone
                        )

                    with st.spinner(
                        f"🔍 Validating {platform} content..."
                    ):

                        validated = validate_content(
                            content,
                            generated,
                            platform
                        )

                    results[platform] = validated

                    progress.progress(
                        (i + 1) / len(platforms)
                    )

                progress.empty()

                # -------- RESULTS --------
                st.success("🎉 Content generated successfully!")

                st.subheader("📱 Platform-Specific Content")

                for platform, result in results.items():

                    with st.expander(
                        f"✨ {platform}",
                        expanded=True
                    ):

                        st.text_area(
                            f"{platform} Content",
                            result,
                            height=350,
                            key=f"{platform}_output"
                        )

                # -------- DOWNLOADS --------
                st.subheader("📥 Export")

                txt_content = create_txt(results)

                docx_file = create_docx(results)

                col1, col2 = st.columns(2)

                with col1:

                    st.download_button(
                        "⬇️ Download TXT",
                        data=txt_content,
                        file_name="repurposed_social_content.txt",
                        mime="text/plain",
                        use_container_width=True
                    )

                with col2:

                    st.download_button(
                        "⬇️ Download DOCX",
                        data=docx_file,
                        file_name="repurposed_social_content.docx",
                        mime=(
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                        use_container_width=True
                    )

            except Exception as e:

                st.error(
                    "Something went wrong while generating the content."
                )

                st.code(str(e))


# ---------------- FOOTER ----------------
st.markdown("---")

st.caption(
    "AI-Based Social Media Content Repurposing System | "
    "Powered by Streamlit + Cloud AI"
)
